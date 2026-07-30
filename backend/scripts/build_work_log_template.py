"""生成不含业务数据的 daily-v1.docx 母版。

只在模板样式需要升级时运行。生成物可以提交 Git，但旧 .doc 样例不能提交。
"""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "templates" / "work_logs" / "daily-v1.docx"
sys.path.insert(0, str(ROOT))

from services.work_log_document import (  # noqa: E402
    BLOCKS,
    _add_native_table,
    _block_rows,
    _definitions,
    _embed_workbook,
    _preview_bytes,
    _workbook_bytes,
)
from services.work_log_schema import (  # noqa: E402
    default_manual_values,
    get_schema,
)


def set_east_asia_font(run, name: str):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def build():
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Pt(104.9)
    section.bottom_margin = Pt(99.2)
    section.left_margin = Pt(76.55)
    section.right_margin = Pt(76.55)
    section.header_distance = Pt(42.55)
    section.footer_distance = Pt(49.6)

    normal = document.styles["Normal"]
    normal.font.name = "方正仿宋_GBK"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
    normal.font.size = Pt(14)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    title_run = title.add_run("滨湖新城派出所社区警务工作日志")
    set_east_asia_font(title_run, "方正小标宋_GBK")
    title_run.font.size = Pt(22)

    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_line.add_run(
        "{{month}}月工作日志（业务日期：{{business_date}}）"
    )
    set_east_asia_font(date_run, "方正仿宋_GBK")
    date_run.font.size = Pt(14)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    note_run = note.add_run("滨湖新城派出所  {{issue_date}}")
    set_east_asia_font(note_run, "方正仿宋_GBK")
    note_run.font.size = Pt(14)

    add_page_number(section.footer.paragraphs[0])
    schema = get_schema()
    definitions = _definitions(schema)
    values = default_manual_values()
    for index, (block_id, block_title, prefixes) in enumerate(
        BLOCKS,
        start=1,
    ):
        if index > 1:
            document.add_section()
        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run(block_title)
        set_east_asia_font(heading_run, "方正小标宋_GBK")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        rows = _block_rows(
            block_id,
            prefixes,
            definitions,
            values,
        )
        _embed_workbook(
            document,
            index=index,
            workbook=_workbook_bytes(block_title, rows),
            preview=_preview_bytes(block_title, rows),
        )
        if block_id == "model_three":
            _add_native_table(
                document,
                "priority.details",
                definitions,
                values,
            )
        elif block_id == "security_venues":
            _add_native_table(
                document,
                "security.details",
                definitions,
                values,
            )
        elif block_id == "security_special":
            _add_native_table(
                document,
                "notices.items",
                definitions,
                values,
            )
            _add_native_table(
                document,
                "special.items",
                definitions,
                values,
            )
    props = document.core_properties
    props.title = "滨湖新城派出所社区警务工作日志"
    props.subject = "脱敏日报母版"
    props.author = "滨湖智慧平台"
    props.last_modified_by = "滨湖智慧平台"
    props.comments = ""
    props.keywords = ""
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
