"""从脱敏母版生成工作日志 DOCX，并内嵌 12 份 XLSX。"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from PIL import Image, ImageDraw, ImageFont


TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent
    / "templates"
    / "work_logs"
    / "daily-v1.docx"
)
XLSX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "spreadsheetml.sheet"
)

# 原日报中的 12 个 Excel 区域在新版中统一为以下 12 块。
BLOCKS = [
    ("basic_population", "基础数据", ["basic."]),
    ("model_three", "重点人员核查", ["priority.model3_"]),
    ("rental_visit", "出租房走访", ["rental.visits", "rental.added", "rental.changed", "rental.cancelled", "rental.total_changes"]),
    ("rental_reverse", "出租房反向核查", ["rental.current_stock", "rental.reverse_checks", "rental.analysis"]),
    ("rental_quality", "出租房质态", ["rental.person_avg_visits", "rental.person_avg_changes", "rental.household_avg_changes", "rental.rated", "rental.rating_rate", "rental.ranking"]),
    ("self_owned_visit", "自购房走访", ["self_owned.visits", "self_owned.added", "self_owned.changed", "self_owned.cancelled", "self_owned.total_changes"]),
    ("self_owned_quality", "自购房质态", ["self_owned."]),
    ("disputes", "矛盾纠纷", ["disputes."]),
    ("fire", "消防", ["fire."]),
    ("security_venues", "行业场所", ["security.venues_checked", "security.hazards"]),
    ("security_dogs", "犬只管理", ["security.dogs"]),
    ("security_special", "黄赌整治与电诈", ["security.special_cases", "security.analysis", "fraud."]),
]


def _font(size: int = 22):
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/msyh.ttc",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def _display_value(value, definition: dict) -> str:
    if value is None or value == "":
        return ""
    if definition["type"] == "percent":
        return f"{float(value):g}%"
    if definition["type"] == "table":
        if not value:
            return "无"
        columns = definition.get("columns") or []
        return "\n".join(
            "；".join(
                f"{column['label']}：{row.get(column['key'], '')}"
                for column in columns
                if row.get(column["key"], "")
            )
            or "空白记录"
            for row in value
        )
    if definition["type"] in {"number", "decimal"}:
        number = float(value)
        return f"{number:g}"
    return str(value)


def _definitions(schema: dict) -> dict[str, dict]:
    return {
        item["id"]: item
        for section in schema["sections"]
        for item in section["fields"]
    }


def _block_rows(
    block_id: str,
    prefixes: list[str],
    definitions: dict[str, dict],
    values: dict,
) -> list[tuple[str, str]]:
    included: list[str] = []
    for field_id in definitions:
        if any(
            field_id == prefix
            or (
                prefix.endswith(".")
                and field_id.startswith(prefix)
            )
            for prefix in prefixes
        ):
            included.append(field_id)

    # 两个拆分块避免重复显示相同字段。
    if block_id == "self_owned_quality":
        included = [
            item for item in included
            if item not in {
                "self_owned.visits", "self_owned.added",
                "self_owned.changed", "self_owned.cancelled",
                "self_owned.total_changes",
            }
        ]
    rows = [
        (
            definitions[field_id]["label"],
            _display_value(values.get(field_id), definitions[field_id]),
        )
        for field_id in included
    ]
    return rows or [("说明", "")]


def _workbook_bytes(title: str, rows: list[tuple[str, str]]) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "工作日志"
    sheet.append([title, ""])
    sheet.merge_cells("A1:B1")
    sheet.append(["项目", "内容"])
    for label, value in rows:
        sheet.append([label, value])
    sheet.column_dimensions["A"].width = 24
    sheet.column_dimensions["B"].width = 72
    thin = Side(style="thin", color="808080")
    for row in sheet.iter_rows(
        min_row=1,
        max_row=sheet.max_row,
        min_col=1,
        max_col=2,
    ):
        for cell in row:
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.alignment = Alignment(
                vertical="center",
                wrap_text=True,
            )
            cell.font = Font(name="宋体", size=10)
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor="D9EAF7")
        cell.font = Font(name="宋体", size=12, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for cell in sheet[2]:
        cell.fill = PatternFill("solid", fgColor="EEF3F8")
        cell.font = Font(name="宋体", size=10, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[1].height = 24
    for index in range(2, sheet.max_row + 1):
        sheet.row_dimensions[index].height = 28
    workbook.calculation.fullCalcOnLoad = True
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _preview_bytes(title: str, rows: list[tuple[str, str]]) -> bytes:
    width = 1240
    left_width = 300
    title_height = 58
    header_height = 48
    line_height = 32
    row_heights = []
    for _label, value in rows:
        row_heights.append(max(52, 20 + line_height * max(1, len(value) // 42 + 1)))
    height = title_height + header_height + sum(row_heights) + 2
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = _font(26)
    body_font = _font(21)
    small_font = _font(19)
    border = "#64748b"
    draw.rectangle((0, 0, width - 1, title_height), fill="#d9eaf7", outline=border)
    draw.text((width / 2, title_height / 2), title, font=title_font, fill="#0f172a", anchor="mm")
    y = title_height
    draw.rectangle((0, y, left_width, y + header_height), fill="#eef3f8", outline=border)
    draw.rectangle((left_width, y, width - 1, y + header_height), fill="#eef3f8", outline=border)
    draw.text((left_width / 2, y + header_height / 2), "项目", font=body_font, fill="#0f172a", anchor="mm")
    draw.text(((left_width + width) / 2, y + header_height / 2), "内容", font=body_font, fill="#0f172a", anchor="mm")
    y += header_height
    for (label, value), row_height in zip(rows, row_heights):
        draw.rectangle((0, y, left_width, y + row_height), outline=border)
        draw.rectangle((left_width, y, width - 1, y + row_height), outline=border)
        draw.text((14, y + 13), label, font=small_font, fill="#0f172a")
        lines = [value[index:index + 42] for index in range(0, len(value), 42)] or [""]
        for line_index, line in enumerate(lines):
            draw.text(
                (left_width + 14, y + 10 + line_index * line_height),
                line,
                font=small_font,
                fill="#0f172a",
            )
        y += row_height
    output = BytesIO()
    image.save(output, format="PNG", optimize=True)
    return output.getvalue()


def _embed_workbook(
    document: Document,
    *,
    index: int,
    workbook: bytes,
    preview: bytes,
):
    part = document.part
    package = part.package
    workbook_part = Part(
        PackURI(f"/word/embeddings/work-log-{index:02d}.xlsx"),
        XLSX_CONTENT_TYPE,
        workbook,
        package,
    )
    package.parts.append(workbook_part)
    workbook_rid = part.relate_to(workbook_part, RT.PACKAGE)
    preview_rid, _ = part.get_or_add_image(BytesIO(preview))
    shape_id = 2048 + index
    object_id = f"_BINHU_WORK_LOG_{index:02d}"
    xml = (
        f'<w:r {nsdecls("w", "r")} '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office">'
        '<w:object w:dxaOrig="8640" w:dyaOrig="4320">'
        f'<v:shape id="_x0000_i{shape_id}" type="#_x0000_t75" '
        'style="width:430pt;height:215pt" o:ole="">'
        f'<v:imagedata r:id="{preview_rid}" o:title=""/>'
        '</v:shape>'
        f'<o:OLEObject Type="Embed" ProgID="Excel.Sheet.12" '
        f'ShapeID="_x0000_i{shape_id}" DrawAspect="Content" '
        f'ObjectID="{object_id}" r:id="{workbook_rid}"/>'
        '</w:object>'
        '</w:r>'
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph._p.append(parse_xml(xml))


def _replace_embedded_workbook(
    document: Document,
    *,
    index: int,
    workbook: bytes,
    preview: bytes,
):
    objects = document.part.element.xpath(
        "//*[local-name()='OLEObject']"
    )
    if len(objects) != len(BLOCKS):
        raise RuntimeError("工作日志母版中的 Excel 对象数量不正确")
    ole_object = objects[index - 1]
    workbook_rid = ole_object.get(qn("r:id"))
    object_container = ole_object.getparent()
    preview_nodes = object_container.xpath(
        ".//*[local-name()='imagedata']"
    )
    if not workbook_rid or len(preview_nodes) != 1:
        raise RuntimeError("工作日志母版中的 Excel 关系不完整")
    preview_rid = preview_nodes[0].get(qn("r:id"))
    workbook_part = document.part.related_parts[workbook_rid]
    preview_part = document.part.related_parts[preview_rid]
    workbook_part._blob = workbook
    preview_part._blob = preview


def _replace_text(document: Document, replacements: dict[str, str]):
    for paragraph in document.paragraphs:
        for marker, value in replacements.items():
            if marker in paragraph.text:
                for run in paragraph.runs:
                    if marker in run.text:
                        run.text = run.text.replace(marker, value)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                for marker, value in replacements.items():
                    if marker in paragraph.text:
                        for run in paragraph.runs:
                            if marker in run.text:
                                run.text = run.text.replace(marker, value)


def _add_native_table(
    document: Document,
    field_id: str,
    definitions: dict[str, dict],
    values: dict,
):
    definition = definitions[field_id]
    columns = definition.get("columns") or []
    rows = values.get(field_id)
    rows = rows if isinstance(rows, list) and rows else [{}]
    heading = document.add_paragraph()
    run = heading.add_run(definition["label"])
    run.bold = True
    set_name = "方正仿宋_GBK"
    run.font.name = set_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), set_name)
    run.font.size = Pt(14)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column["label"]
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row.get(column["key"], "") or "")
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for cell_run in paragraph.runs:
                    cell_run.font.name = "宋体"
                    cell_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    cell_run.font.size = Pt(9)


def _replace_native_table(
    document: Document,
    *,
    index: int,
    field_id: str,
    definitions: dict[str, dict],
    values: dict,
):
    if len(document.tables) != 4:
        raise RuntimeError("工作日志母版中的普通表格数量不正确")
    table = document.tables[index]
    definition = definitions[field_id]
    columns = definition.get("columns") or []
    rows = values.get(field_id)
    rows = rows if isinstance(rows, list) and rows else [{}]
    for existing in list(table.rows[1:]):
        table._tbl.remove(existing._tr)
    for row in rows:
        cells = table.add_row().cells
        for column_index, column in enumerate(columns):
            cells[column_index].text = str(
                row.get(column["key"], "") or ""
            )
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for cell_run in paragraph.runs:
                    cell_run.font.name = "宋体"
                    cell_run._element.rPr.rFonts.set(
                        qn("w:eastAsia"),
                        "宋体",
                    )
                    cell_run.font.size = Pt(9)
def build_daily_document(
    draft: dict,
    schema: dict,
    values: dict,
) -> tuple[bytes, str]:
    if not TEMPLATE_PATH.exists():
        raise RuntimeError("工作日志母版缺失")
    document = Document(TEMPLATE_PATH)
    snapshot = draft["system_snapshot"]
    business_date = snapshot["business_date"]
    issue_date = snapshot["issue_date"]
    _replace_text(
        document,
        {
            "{{business_date}}": business_date,
            "{{issue_date}}": issue_date,
            "{{month}}": str(snapshot["month"]),
        },
    )
    definitions = _definitions(schema)
    for index, (block_id, title, prefixes) in enumerate(BLOCKS, start=1):
        rows = _block_rows(block_id, prefixes, definitions, values)
        _replace_embedded_workbook(
            document,
            index=index,
            workbook=_workbook_bytes(title, rows),
            preview=_preview_bytes(title, rows),
        )
    for index, field_id in enumerate((
        "priority.details",
        "security.details",
        "notices.items",
        "special.items",
    )):
        _replace_native_table(
            document,
            index=index,
            field_id=field_id,
            definitions=definitions,
            values=values,
        )

    properties = document.core_properties
    properties.title = "滨湖新城派出所社区警务工作日志"
    properties.subject = "工作日志"
    properties.author = "滨湖智慧平台"
    properties.last_modified_by = "滨湖智慧平台"
    properties.comments = ""
    properties.keywords = ""
    output = BytesIO()
    document.save(output)
    filename = (
        f"{snapshot['filename_prefix']}"
        "日报滨湖新城派出所社区警务工作日志.docx"
    )
    return output.getvalue(), filename
